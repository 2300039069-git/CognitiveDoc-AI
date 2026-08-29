import os
import re
from pathlib import Path
from typing import Dict, Any, Tuple
import pypdf
import docx

def clean_extracted_text(text: str) -> str:
    """Normalize whitespace and remove unprintable characters."""
    if not text:
        return ""
    # Normalize multiple newlines and spaces
    text = re.sub(r'\r\n|\r', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def extract_from_pdf(file_path: Path) -> Tuple[str, int]:
    """Extract text and page count from PDF using high-speed C++ PyMuPDF (fitz)."""
    text_parts = []
    page_count = 0
    try:
        import fitz
        doc = fitz.open(str(file_path))
        page_count = len(doc)
        for i in range(page_count):
            page_text = doc[i].get_text("text") or ""
            if page_text.strip():
                text_parts.append(f"--- Page {i+1} ---\n{page_text}")
        doc.close()
    except Exception:
        try:
            reader = pypdf.PdfReader(str(file_path))
            page_count = len(reader.pages)
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(f"--- Page {i+1} ---\n" + page_text)
        except Exception as e:
            text_parts.append(f"Error extracting PDF: {str(e)}")
            page_count = max(page_count, 1)

    full_text = clean_extracted_text("\n\n".join(text_parts))
    return full_text, max(page_count, 1)

def extract_from_docx(file_path: Path) -> Tuple[str, int]:
    """Extract text from DOCX file."""
    text_parts = []
    try:
        doc = docx.Document(str(file_path))
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    text_parts.append(" | ".join(row_cells))
    except Exception as e:
        text_parts.append(f"Error extracting DOCX: {str(e)}")

    full_text = clean_extracted_text("\n\n".join(text_parts))
    # Approximate page count (~400 words per page)
    word_count = len(full_text.split())
    page_count = max(1, (word_count + 399) // 400)
    return full_text, page_count

def extract_from_txt(file_path: Path) -> Tuple[str, int]:
    """Extract text from TXT or Markdown file."""
    text = ""
    for enc in ['utf-8', 'latin-1', 'cp1252']:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                text = f.read()
                break
        except Exception:
            continue
    full_text = clean_extracted_text(text)
    word_count = len(full_text.split())
    page_count = max(1, (word_count + 399) // 400)
    return full_text, page_count

def extract_text_and_metadata(file_path: Path, original_filename: str) -> Dict[str, Any]:
    """Main entrypoint for document extraction."""
    ext = Path(original_filename).suffix.lower()
    file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

    if ext == ".pdf":
        text, pages = extract_from_pdf(file_path)
        file_type = "pdf"
    elif ext in [".docx", ".doc"]:
        text, pages = extract_from_docx(file_path)
        file_type = "docx"
    elif ext in [".txt", ".md", ".csv", ".log", ".json"]:
        text, pages = extract_from_txt(file_path)
        file_type = ext.replace(".", "")
    else:
        text, pages = extract_from_txt(file_path)
        file_type = "txt"

    words = [w for w in text.split() if w]
    word_count = len(words)
    char_count = len(text)
    estimated_read_time = round(word_count / 200, 1)  # 200 wpm average reading speed

    return {
        "text": text,
        "file_type": file_type,
        "file_size": file_size,
        "page_count": pages,
        "word_count": word_count,
        "char_count": char_count,
        "read_time_minutes": estimated_read_time
    }
